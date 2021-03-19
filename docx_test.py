#前チェック用モジュール

import sqlite3
import datetime
import docx
import os
from tkinter import messagebox

#クエリテストの実装
def ins_test(conn):
    cur = conn.cursor()
    query = 'select count(*) from persons'

    #値をselect
    if len(cur.fetchall()) == 0:
        dat = cur.execute(query)
    print(cur.fetchall())
    return cur.fetchall()

def sel_test():
    dbname = 'soyokaze.db'
    conn = sqlite3.connect(dbname)
    sel_query(conn)
    print(sel_query(conn))  
    conn.close()

#クエリ実体
##db作成
def crete_tbl(conn):
    cur = conn.cursor()
    cur.execute(
    'CREATE TABLE IF NOT EXISTS persons(id INTEGER PRIMARY KEY AUTOINCREMENT,moditime \
    ,syubetu1,syubetu2,name1,sex,name2,furi2,birthday,startday,endday,address,phonenumber \
    ,shiharaihou,shiharaiyen,roudoutime,stattime,endtime \
    ,koutuyen,koutukm,nenkinno,hihokenno,hihokenother \
    ,haiguumu,haiguyen,hihoen_ex)')

##Sel
def sel_query(conn):
    cur = conn.cursor()
    query = 'select * from persons'
    dat = cur.execute(query)
    return cur.fetchall()

##count その1 そもそもレコードが0件の場合はチェックしない
def count_query0(conn):
    cur = conn.cursor()
    query = 'select * from persons'
    dat = cur.execute(query)
    return len(cur.fetchall())

##count その2 テストの本体内容
def countest_query1(conn,key0,val0,key1,val1,key2,val2):
    cur = conn.cursor()
    query = 'select * from persons where 1 = 1'
    query = query + ' and ' +  str(key0) + '= '  +  '\'' + str(val0) + '\''
    query = query + ' and ' +  str(key1) + '= '  +  '\'' + str(val1) + '\''
    query = query + ' and ' +  str(key2) + '= '  +  '\'' + str(val2) + '\''
    
    print(query)  
    dat = cur.execute(query)
    #値をselect
    return len(cur.fetchall())

#insert チェック済データをPKと時間つきで投入する
def ins_beforedata(conn,vallist):
    cur = conn.cursor()

    dt_now = datetime.datetime.now()
    que = '\'' + str(dt_now) + '\'' + ','  

    i = 0
    #取り込んだデータをコミット
    for item in vallist:
        if i != 0:
            que = que + ','
        que = que + '\'' + str(item) + '\''
        i = i + 1

    print(que)

    que = 'INSERT INTO persons (moditime,syubetu1,syubetu2,name1,sex,name2,furi2,birthday,startday,endday,address,phonenumber \
    ,shiharaihou,shiharaiyen,roudoutime,stattime,endtime \
    ,koutuyen,koutukm,nenkinno,hihokenno,hihokenother \
    ,haiguumu,haiguyen,hihoen_ex) VALUES(' + que + ')'
    print(que)
    cur.execute(que) 

    conn.commit()
    conn.close()
    

##呼び出し本体
def before_check(txt):
    doc = docx.Document(txt.get())
   
    tbl1 = doc.tables[0]
    tbl2 = doc.tables[1]

    total = 0
    total = len(tbl1.rows) + len(tbl2.rows)    

    #比較用に必要なデータをすべて取り込む
    vallist = list(range(total))
    namelist = list(range(total))
    
    idx = 0

    for row in tbl1.rows:
        cells1 = row.cells
        vallist[idx] = cells1[3].text
        namelist[idx] = cells1[0].text + cells1[1].text
        print(namelist[idx])
        print(vallist[idx])
        idx = idx + 1

    for row in tbl2.rows:
        cells2 = row.cells 
        vallist[idx] = cells2[3].text
        namelist[idx] = cells1[0].text + cells1[1].text
        print(namelist[idx])
        print(vallist[idx])
        idx = idx + 1

    dbname = 'soyokaze.db'
    conn = sqlite3.connect(dbname)
    crete_tbl(conn)

    #取り込みフォーマット成型
    vallist[18] = str.upper(vallist[18])
    vallist[19] = str.upper(vallist[19])

    #レコードがない場合は挿入して終わり
    #レコードがある場合はテストして問題なければ挿入
    if 0 == count_query0(conn):
        ins_beforedata(conn,vallist)
        return '同じデータは無かったので、データを登録しました'
    else:
        #同一の手続き種別で、氏名、生年月日、フリガナ、適用開始終了、基礎年金番号
        #被保険者番号が一致の場合はエラーとする
        errtext = ''
        if vallist[1] == '両方' or vallist[1] == '健保年金':
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','健保年金','name2',vallist[4]):
                errtext = errtext + '名前が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','健保年金','furi2',vallist[5]):
                errtext = errtext + 'フリガナが同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','健保年金','birthday',vallist[6]):
                errtext = errtext + '生年月日が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','健保年金','nenkinno',vallist[18]):
                errtext = errtext + '基礎年金番号が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','健保年金','hihokenno',vallist[19]):
                errtext = errtext + '被保険者番号が同じデータが過去にあります\n'

        if vallist[1] == '両方' or vallist[1] == '雇保':
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','雇保','name2',vallist[4]):
                errtext = errtext + '名前が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','雇保','furi2',vallist[5]):
                errtext = errtext + 'フリガナが同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','雇保','birthday',vallist[6]):
                errtext = errtext + '生年月日が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','雇保','nenkinno',vallist[18]):
                errtext = errtext + '基礎年金番号が同じデータが過去にあります\n'
            if 0 < countest_query1(conn,'syubetu1',vallist[0],'syubetu2','雇保','hihokenno',vallist[19]):
                errtext = errtext + '被保険者番号が同じデータが過去にあります\n'

        if vallist[1] != '両方' and vallist[1] != '雇保' and vallist[1] != '健保年金':
            errtext = errtext = '手続き種別は「健保保険」か「雇保」か「両方」にして下さい '
       
        print(errtext)

        # メッセージボックス（はい・いいえ） 
        #ret = messagebox.askyesno('強制的に登録しますか？', errtext)
        #if ret == True:
        #    ins_beforedata(conn,vallist)
        #return errtext
